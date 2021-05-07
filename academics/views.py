
from django.shortcuts import render,redirect
from django.db.models import Subquery
import docx
from django.utils.safestring import SafeString
from django.contrib.auth.models import User , auth , Group
import datetime as dt
from .models import *
import random
from django.conf import settings

import os
import json
from django.core.files.storage import FileSystemStorage
from django.core.paginator import Paginator
from django.core.mail import send_mail , EmailMessage


from django.contrib import messages


def dashboardView(request):
    print(request.user.is_authenticated)
    if request.user.is_authenticated:
        dataObj = User.objects.get(username=request.user).groups.all()
        context={}
        for i in dataObj:
            context[str(i.name)]=True
        print(context)
        return render(request,'dashboard.html',context)
    else:
        return redirect('/accounts/login')




#------------------------ MAIN ACADEMICS --------------------------------------------
def Etype(request,id):
    print(id)
    id = employee.objects.filter(EmployeeID=request.user)
    print(id)
    for i in id:
        id = i.Etype
    print("employee",id)
    if stud_details.objects.filter(UniversityEmailID=request.user).exists():
        return redirect('/socrates')
    elif employee.objects.filter(EmployeeID = request.user).exists():
        
        if id == "DEANSOE":
            return render(request,'dean/dashboard.html')
        elif id =="DEANSOM":
            return render(request,'dean/dashboard.html')
        elif id =="COE":
            return render(request,'coe/dashboard.html')
        elif id== "HOD":
            return render(request,'hod/dashboard.html')
        elif id=="ASSISTANT-PROF":
            return render(request,'professors/dashboard.html')
        elif id=="REGISTRAR":
            return render(request,'registrar/dashboard.html')
        elif id == "VICE-CHANCELLOR":
            return render(request,'vc/dashboard.html')
    else:
        return HttpResponse("ERROR 404 File not Found..!!")



def lecture(request):
    print(request.user)
    id = employee.objects.filter(EmployeeID = request.user)
    for i in id:
        id = i.Department
    print(id)
    students = stud_details.objects.filter(Department=id).count
    faculti = faculty.objects.filter(Department =id).count
    lecture = lectures.objects.filter(Department = id)
    context ={
        'department':department,
        'students' : students,
        'faculty' :faculti,
        'lectures' : lecture,
        'lecture_count' : lecture.count()
    }
    courses = faculty.objects.all().order_by('FieldOfExpertise').distinct()
    print(courses)
    return render(request , 'hod/lecture.html',context)

def studentLectureView(request):
    lectures = lectureEnrollment.objects.filter(student = stud_details.objects.get(UniversityEmailID=request.user))
    return render(request,'student/lecture.html',{'lecture':lectures})


def studentTrackView(request,lecture):
    obj = lectureRecord.objects.filter(lectureId=lectures.objects.get(id=lecture))
    return render(request,'bs4_Horizontal_timeline.html',{'track':obj,'student':True})

def AddLecture(request):
    if request.method=="POST":
        sub = request.POST['SubName']
        sem = request.POST['Sem']
        elig_teacher = request.POST['teacher']
        credit = request.POST['credit']
        branch = request.POST['Branch']
        while True:
            code = branch +"-"+ str(sem)+ str(random.randint(1,99))
            if lectures.objects.filter(LectureID = code).exists():
                continue
            else:
                break
        print(code)
        lec = lectures(LectureID = code,
                LectureName = sub,
                Semester = sem,
                TaughtBy = elig_teacher,
                Credit = credit)
        lec.save()

        return redirect('/academic/hod')

def remove(request):
    if request.method=="POST":
        code = request.POST['code']
        if lectures.objects.filter(LectureID = code).exists():
            lectures.objects.filter(LectureID = code).delete()
            return redirect('/academic/lecture/view')
        else:
            return redirect('/academic/lecture/view')
        
def faculty_info(request):
    faculties = faculty.objects.all()
    return render(request,'dean/department.html',{'faculties':faculties})
    
def deanDepartmentView(request):
    course = Department.objects.filter(School=school.objects.get(dean=employee.objects.get(Email=request.user)))
    print(course)
    return render(request,'dean/department.html',{'courses':course})

def deanDepartmentDetailView(request,id):
    department = Department.objects.filter(DepartmentID = id)
    students = stud_details.objects.filter(Department=id).count
    faculti = faculty.objects.filter(Department =id).count
    lecture = lectures.objects.filter(Department = id)
    context ={
        'department':department,
        'students' : students,
        'faculty' :faculti,
        'lectures' : lecture,
        'lecture_count' : lecture.count()
    }
    return render(request,'dean/lecture.html',context)

    # elif id=="add":
    #     if request.method=="POST":
    #         Deptname = request.POST['DeptName']
    #         DeptID = request.POST['DeptID']
    #         head = request.POST['head']
    #         schoolData = request.POST['School']
    #         if Department.objects.filter(DepartmentID = DeptID ).exists() or Department.objects.filter(DepartmentName=Deptname).exists() or Department.objects.filter(HeadOfDepartment=head).exists():  
    #             messages.info(request,'Invalid Credentials')
    #             return redirect('/academic/department/add')
    #         else:
    #             obj = Department(
    #                 DepartmentName=Deptname , DepartmentID = DeptID , HeadOfDepartment='none' , School=school
    #             , StartDate = dt.datetime.now().date(),
    #             SpecializedCourse=False,Status = False)
    #             obj.save()
    #             subject = 'Assigning new username of hod'
    #             message = 'Hello , we need a new username of the category HOD for the'+ str(Deptname)+' and assign it to the' + (head)+"\nThanks,\n DEAN SOE"
    #             email = EmailMessage(subject, 
    #                 message, EMAIL_HOST_USER, ['aniket.vyas@spsu.ac.in'])
    #             email.fail_silently = False
    #             email.send()
                
    #             print("department can be created")
    #             messages.info(request,'Department Added Successfully')
    #             return redirect('/academic/dean')
    #     heads=faculty.objects.all().order_by('Department')
    #     return render(request,'dean/department.html',{'heads':heads,'addDepartment':True})
    # else:
    #     department = Department.objects.filter(DepartmentID = id)
    #     students = stud_details.objects.filter(Department=id).count
    #     faculti = faculty.objects.filter(Department =id).count
    #     lecture = lectures.objects.filter(Department = id)
       
    #     context ={
    #         'department':department,
    #         'students' : students,
    #         'faculty' :faculti,
    #         'lectures' : lecture,
    #         'lecture_count' : lecture.count(),
    #         'heads': heads
    #     }
    #     return render(request,'dean/department.html',context)

def DepartmentFaculty(request,id):
    return render(request,'dean/department.html',{'faculties':faculty.objects.all()})

#------------------------------ASSISTANT PROFESSORS -----------------------------

def asstlecture(request):
    obj = lectures.objects.filter(TaughtBy=faculty.objects.get(Email = request.user))
    return render(request,'professors/lecture.html',{'lect':obj,'lecture':True})

def assLectureView(request,lecture):
    lectur = lectureEnrollment.objects.filter(lecture=lectures.objects.get(pk=lecture))
    lects = lectures.objects.filter(pk=lecture)
    return render(request,'professors/lecture.html',{'outs':lectur,'ongoing':lects})

def lectureRequest(request):
    if request.method == "POST":
        a = request.FILES['lesson']
        b = request.POST['credit']
        c = request.POST['subjectName']
        d = request.POST['semester']
        t = faculty.objects.get(FacultyID=request.user).Department
        print(t)
        id =  random.randint(1,1000)
        filename = fs.save(str(t)+"/"+str(id)+'/'+str(c) , a)
        image_url = fs.url(filename)
        print(image_url)
        lect = TempLectures(
            TempLectureID =id,
            TempLessonPlan = image_url,
            CreatedBy = request.user.username,
            Department = t,
            Semester = d,
            LectureTitle = c,
            LastApprovedBy="PROFESSOR",
            LectureStatus = "UC",
            LastChangesMadeBy = "PROFESSOR",
            NextSendTo = "HOD",
            Changes = "NONE"
        )
        lect.save()
        return redirect('/academic/ASSTPROF/lecture/view')
    else:
        return render(request,'professors/lecture.html',{'lectureRequest':True})  

    
def lectureStatus(request):
    lect = TempLectures.objects.filter(CreatedBy = request.user.username,NextSendTo='PROFESSOR')
    return render(request,'professors/lecture.html',{'lect':lect,'lect1':lect,'lectureStatus':True,'trackedLecture':True})

def lectureChangeInfo(request,id):
    if request.method=="POST":
        a = request.FILES['lesson']
        credit = request.POST['credit']
        subject = request.POST['subject']
        sem = request.POST['semester']
        t = faculty.objects.filter(FacultyID=request.user)
        print(t)
        for i in t:
            t = str(i.Department)
        print(t)
        filename = fs.save(t+"/"+str(subject)+"/"+a.name , a)
        image_url = fs.url(filename)
        print(image_url)
        TempLectures.objects.filter(TempLectureID = id).update(
            Credit=credit,
            LectureTitle =subject , 
            TempLessonPlan=image_url,
            Semester = sem,
            LastChangesMadeBy="PROFESSOR",
            NextSendTo="HOD")
        return redirect('/academic/ASSTPROF/lecture/view')



def trackLecture(request,id):
    if request.method == "POST":
        startDate = request.POST['startDate']
        endDate = request.POST['endDate']
        goals = request.POST['goals']
        if lectureRecord.objects.filter(startDate=startDate,endDate=endDate):
            messages.error(request,'Goals for today already exists')
            return redirect(f'/academic/ASSTPROF/lecture/{id}/Track',id)
        else:
            record = lectureRecord(startDate=startDate,endDate=endDate,goals=goals,
            lectureId = lectures.objects.get(pk=id),
            facultyId = faculty.objects.get(Email=request.user)).save()
            return redirect(f'/academic/ASSTPROF/lecture/{id}/view',id)
    else:
        print(faculty.objects.filter(FacultyID=request.user),lectures.objects.filter(pk=id))
        obj = lectureRecord.objects.filter(facultyId=faculty.objects.get(Email=request.user),lectureId=lectures.objects.get(pk=id))
        return render(request,'bs4_Horizontal_timeline.html',{'trackLecture':True,'trackedLecture':True,"track":obj,'prof':True})

#--------------------------- Attendence ----------------------------------


def ProfAttendence(request,id):
    if request.method=="POST":
        batch_list = request.POST.getlist('batch')
        dateOfAttendance = request.POST['dateOfAttendence']
        if dateOfAttendance:
            print(dateOfAttendance)
            datee = dt.datetime.strptime(dateOfAttendance,"%Y-%m-%d")
            for i in range(0,len(batch_list)):
                batch_list[i]= int(batch_list[i])
            print(stud_details.objects.filter(id__in = batch_list))
            if lectures.objects.filter(pk =id).exists():
                b= lectures.objects.get(pk=id)
                print(datee.date())
                if datee.date() == dt.datetime.now().date():
                    if attendence.objects.filter(LectureID=id).exists() :   # IF Attendance Exists Then
                        a = attendence.objects.get(LectureID=id).attendanceFile
                        a=str(a)
                        print("media/"+a)
                        f = open("media/"+a,'r')
                        y = json.load(f)
                        f.close()
                        print(y['Attendance'],dateOfAttendance)
                        if dateOfAttendance in y["Attendance"]:
                            messages.error(request,'Attendance Already exists..!')
                        else:   # Else Create a new json file to store the Attendance Data
                            f = open("media/"+a,'w')
                            print(y["Attendance"])
                            y["Attendance"][dateOfAttendance]=batch_list
                            json.dump(y,f,indent=6)
                            f.close()
                            messages.info(request,'Attendance uploaded..!!')
                    else:
                        f = open('media/attendance/'+str(b.LectureName)+"_"+str(dt.datetime.now().year)+".json",'w')
                        y={"Attendance":{ dateOfAttendance : batch_list}}
                        json.dump(y,f,indent=6)
                        print("file saved..!!")
                        attendence(
                            LastUpdatedOn=dt.datetime.now().date(),
                            FacultyID =faculty.objects.get(FacultyID=request.user),
                            attendanceFile='attendance/'+str(b.LectureName)+"_"+str(dt.datetime.now().year)+".json",
                            LectureID = b
                            ).save()
                else:
                    messages.error(request,'Attendance can be uploaded of today..!! i.e.'+str(dt.datetime.now().date()))
            else:
                a = attendence(FacultyID = faculty.objects.get(FacultyID=request.user),
                LectureID = lectures.objects.get(pk =id),
                LastUpdatedOn = dt.datetime.now().date(),
                )
                a.save()   
            for i in batch_list:
                data = stud_details.objects.get(id=i)
                #print(data.FirstName)
                #attendenceDetails(AttendenceId=a , StudentID = data, AttendenceDate = dateofattendence).save()
            return redirect('/academic/ASSTPROF/lecture/'+str(id)+'/attendence')    
        else:
            messages.error(request,'Enter Date OF Attendance')
            return redirect('/academic/ASSTPROF/lecture/'+str(id)+'/attendence')
    else:        
        #attendies = lectureEnrollment.objects.filter(lectureId=lectures.objects.get(TaughtBy=faculty.objects.get(FacultyID=request.user)))
        attendies = stud_details.objects.filter(id__in =Subquery(lectureEnrollment.objects.filter(lectureId=lectures.objects.get(pk=id)).values('studentID')))
        print(lectureEnrollment.objects.filter(lectureId=lectures.objects.get(pk=id)),id)
        return render(request,'professors/attendence.html',{'attendies':attendies})


def ProfAttendenceView(request,id):
    text = request.GET.get('button-text')
    if True:
        a=attendence.objects.get(LectureID=(lectures.objects.get(pk=id))).attendanceFile
        print("aaaa",a)
        f = open("media/"+str(a),'r')
        y=json.load(f)
        f.close()
        print(y)
        arr= y['Attendance']
        print(y['Attendance'])
        studentDictionary={}
        name ,i,j=0,0,0
        for i in arr.keys():    # i contains the list of present students
            k=arr[i]
            print(k)
            for j in range(0,len(k)):
                name = stud_details.objects.get(id=k[j])
                name = name.FirstName +" " +name.LastName
                if name in studentDictionary.keys():
                    studentDictionary[name]+=1
                else:
                    studentDictionary[name]=1
        main_list=[]
        print(arr.keys())
        total =len(arr.keys())
        for i in studentDictionary.keys():
            k = studentDictionary[i]
            l={"name":i,"percentage":(k/total)*100}
            main_list.append(l)
        print(main_list)

        return render(request,'professors/attendence.html',{'students':main_list})
        #return HttpResponse(json.dumps(main_list))
        #return JsonResponse(data) 




#---------------------------------Head Of Department--------------------------------
def hodLecture(request):
    lect = lectures.objects.filter(Department = Department.objects.get(HeadOfDepartment = employee.objects.get(Email=request.user)))
    return render(request,'hod/lecture.html',{'lectureInformation':lect,'lectureInfo':True})

def HodLectureRequest(request):
    lect = TempLectures.objects.filter(NextSendTo = "HOD")
    return render(request,'hod/lecture.html',{'lect':lect})


def lectureConfigureRequest(request,id):
    user = employee.objects.get(EmployeeID=request.user).Etype
    if user == "HOD":
        TempLectures.objects.filter(TempLectureID=id).update(NextSendTo="PROFESSOR")
        return redirect('/academic/HOD/lecture/lectureRequest')
    elif user == 'DEANSOE':
        TempLectures.objects.filter(TempLectureID=id).update(NextSendTo="HOD")
        return redirect('/academic/deanEng/lecture/lectureRequest')


# def OngoingLectureView(request):
#     lect = lectures.objects.filter(Department=Department.objects.get(HeadOfDepartment = request.user.username))
#     return render(request,'hod/lectureOngoing.html',{'lectures':lect})


def LectureEnrollments(request,id):
    l = stud_details.objects.filter(Department = Department.objects.get(HeadOfDepartment=employee.objects.get(Email=request.user))).exclude(id__in = Subquery(lectureEnrollment.objects.filter(lecture=lectures.objects.get(id=id)).values('student')))
    context = {
        'student' : l,
        'lectureEnroll':True
    }
    print(l)
    return render(request,'hod/lecture.html',context)

def LectureEnrollmentDone(request,LectureID,EnrollmentNumber):
    data = stud_details.objects.get(EnrollmentNumber = EnrollmentNumber)
    lect = lectures.objects.get(id = LectureID)
    lectureEnrollment(
        student = data,
        lecture = lect
    ).save()
    return redirect('/academic/HOD/lecture/'+str(LectureID)+'/Enrollment')

def lectureFilter(request , id):
    if request.method == "POST":
        sem = int(request.POST['sem'])
        l = stud_details.objects.filter(CurrSem=sem,Department = Department.objects.get(HeadOfDepartment=employee.objects.get(Email=request.user))).exclude(id__in = Subquery(lectureEnrollment.objects.filter(lectureId=lectures.objects.get(id=id)).values('studentID')))
        print(l)
        return render(request, 'hod/lecture.html',{'student':l,'lectureEnroll':True})


def lol(request,lecture):
    lectur = lectureEnrollment.objects.filter(lecture=lectures.objects.get(id=lecture))
    lects = lectures.objects.filter(id=lecture)
    return render(request,'hod/lecture.html',{'outs':lectur,'ongoing':lects})

def trackLectureHod(request,id):
    obj = lectureRecord.objects.filter(lectureId=lectures.objects.get(id=id))
    return render(request,'bs4_Horizontal_timeline.html',{'track': obj,'hod':True })





def OutpassRequests(request):
    #ik@123
    #p = stud_outing.objects.raw('select "Department_id" from academics_stud_details as sd , outpass_stud_outing as so where sd."EnrollmentNumber" = so."EnrollmentNumber" ;')
    e = Department.objects.get(HeadOfDepartment=request.user.username).DepartmentID
    print(e)
    a = stud_outing.objects.filter(EnrollmentNumber__in=Subquery(stud_details.objects.filter(Department=e).values('EnrollmentNumber')),Req_status='Forwaded')
    return render(request,'hod/outpass.html',{'out':a})

    
#-------------------------DEAN OF ENGINEERING ------------------------------------------

def deanLecture(request):
    return render(request,'dean/lecture.html')

def DeanLectureRequest(request):
    lect = TempLectures.objects.filter(NextSendTo = "DEAN")
    return render(request,'dean/lecture.html',{'lect':lect})

def DeanOngoingRequest(request):
    lect = lectures.objects.all()
    return render(request,'dean/lecture.html',{'lect':lect})
    
def lectureChangeRequest(request,id):
    print(id)
    asl = id
    if request.method=="POST":
        content = request.POST['Changes']
        branch = TempLectures.objects.filter(TempLectureID=id)
        for i in branch:
            Department = i.Department
            subject = i.LectureTitle
            change = i.Changes
            fil = i.TempLectureID
            subject = i.LectureTitle
        print(Department,content)
        if change == 'NONE':
            document = Document()
            print(request.user.get_full_name)
            document.add_heading('\tChanges Suggested By '+str(request.user.get_full_name())+' on'+str(dt.datetime.now().date())+'\t')
            document.add_paragraph(content)
            document.save('media/'+str(Department)+"/"+str(asl)+"/"+str(id)+'.docx')
            change='media/'+str(Department)+"/"+str(asl)+"/"+str(id)+'.docx'
        else:
            print("change",change)
            document = Document('media'+"/"+str(Department)+"/"+str(asl)+"/"+str(id)+'.docx')
            print('docuemet created')
            print(request.user.get_full_name)
            document.add_heading('\tChanges Suggested By '+str(request.user.get_full_name())+'\t')
            document.add_paragraph(content)
            print(change)
            document.save('media'+"/"+str(Department)+"/"+str(asl)+"/"+str(id)+'.docx')
        print("File created Successfully")

        TempLectures.objects.filter(TempLectureID=id).update(Changes=change)
        id = employee.objects.filter(EmployeeID=request.user)
        for i in id:
            id = i.Etype
        print(id)
        if id=='HOD':
            a=TempLectures.objects.filter(TempLectureID=asl)
            for i in a:
                print(i.LastChangesMadeBy,i.NextSendTo)
            TempLectures.objects.filter(TempLectureID=asl).update(LastChangesMadeBy='HOD',NextSendTo='PROFESSOR')
            a=TempLectures.objects.filter(TempLectureID=asl)
            for i in a:
                print(i.LastChangesMadeBy,i.NextSendTo)
            lect = TempLectures.objects.filter(LectureStatus='UC',NextSendTo='HOD')
            return redirect('/academic/HOD/lecture/lectureRequest',{'lect':lect})
        elif id == "DEANSOE":
            TempLectures.objects.filter(TempLectureID=asl).update(LastChangesMadeBy='DEAN',NextSendTo='HOD')
            lect = TempLectures.objects.filter(LectureStatus='UC',NextSendTo='DEAN')
            return redirect('/academic/deanEng/lecture/lectureRequest',{'lect':lect})

        elif id == "VICE-CHANCELLOR":
            TempLectures.objects.filter(TempLectureID=asl).update(LastChangesMadeBy='VC',NextSendTo='DEAN')
            lect = TempLectures.objects.filter(LectureStatus='UC',NextSendTo='VC')
            return redirect('/academic/vc/lecture/lectureRequest',{'lect':lect})

        

def lectureApproveRequest(request,id):
    user = request.user
    userType = employee.objects.filter(EmployeeID=user)
    for i in userType:
        userType = i.Etype
        de = i.Department
    print(userType)
    if userType=="HOD":
        lect = TempLectures.objects.filter(Department=de)
        TempLectures.objects.filter(TempLectureID=id).update(LastApprovedBy='HOD',NextSendTo='DEAN')
        return redirect('/academic/HOD/lecture/lectureRequest',{'lect':lect})
    elif userType=="DEANSOE":
        lect = TempLectures.objects.filter(Department=de)
        TempLectures.objects.filter(TempLectureID=id).update(LastApprovedBy='DEAN',NextSendTo="VC")
        return redirect('/academic/deanEng/lecture/lectureRequest',{'lect':lect})



#-------------------------- VICE - CHANCELLOR ----------------------------------
def lectureView(request):
    department = Department.objects.all().count
    students = stud_details.objects.all().count
    faculti = faculty.objects.all().count
    lecture = lectures.objects.all()
    context ={
            'department':department,
            'students' : students,
            'faculty' :faculti,
            'lectures' : lecture,
            'lecture_count' : lecture.count()
    }
    return render(request,'vc/lectures.html',context)


def vclectureRequest(request): 
    lect = TempLectures.objects.filter(NextSendTo="VC")
    return render(request,'vc/lectureStatus.html',{'lect':lect})


def lectureApproveVc(request,id):
    temp = TempLectures.objects.get(TempLectureID=id)
    while True:
        id = random.randint(10000,999999)
        if lectures.objects.filter(id=id).exists():
            continue
        else:
            break
    lect = lectures(
        id = id,
        LectureID = temp.TempLectureID,
        LectureName = temp.LectureTitle,
        LessonPlan = temp.TempLessonPlan,
        TaughtBy = faculty.objects.get(FacultyID = temp.CreatedBy),
        Credit = temp.Credit,
        LastTaughtBy = "No",
        Department = temp.Department,
        Semester = temp.Semester
    )
    lect.save();
    temp.delete()
    return redirect('/academic/vc')






